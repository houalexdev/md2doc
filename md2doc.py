#!/usr/bin/env python3
# ===== md2doc.py (主程序) =====

import json
import sys
import warnings
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from styles_loader import load_styles
from numbering import create_numbering_definition
from md_parser import parse_md
from docx_builder import (
    add_codeblock, add_heading, add_list_item, 
    add_paragraph, add_table
)

warnings.filterwarnings("ignore", category=UserWarning, module="docx.styles.styles")

STYLES_FILE = "styles.json"

def md_to_docx(md_text: str, out_file: str):
    """主转换函数：从 Markdown 转换为 DOCX。"""
    cfg = json.load(open(STYLES_FILE, "r", encoding="utf-8"))
    doc = Document()

    # 默认中文字体
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    if hasattr(normal, "_element"):
        rPr = normal._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "宋体")

    load_styles(doc, cfg)
    
    # 创建编号定义
    current_num_id = create_numbering_definition(doc, cfg)

    for kind, *args in parse_md(md_text):
        if kind == "codeblock":
            lang, content = args
            add_codeblock(doc, lang, content)
            
        elif kind == "heading":
            lvl, txt = args
            add_heading(doc, lvl, txt)
            # 每次遇到标题，创建新的编号实例
            current_num_id = create_numbering_definition(doc, cfg)
            
        elif kind == "list_item":
            level, txt = args
            add_list_item(doc, level, txt.lstrip(), current_num_id)
            
        elif kind == "para":
            txt = args[0]
            add_paragraph(doc, txt)
            
        elif kind == "table":
            rows = args[0]
            add_table(doc, rows, cfg)

    doc.save(out_file)
    print("→ Saved", out_file)


if __name__ == "__main__":
    in_file = sys.argv[1] if len(sys.argv) > 1 else "input.md"
    out_file = sys.argv[2] if len(sys.argv) > 2 else "output.docx"

    md_text = open(in_file, encoding="utf-8").read()
    md_to_docx(md_text, out_file)