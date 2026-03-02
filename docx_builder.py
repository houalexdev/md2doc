#!/usr/bin/env python3
# ===== docx_builder.py =====
"""DOCX 文档生成模块"""
import re
import sys
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import rgb
from numbering import apply_list_numbering, create_numbering_definition

def apply_bold_runs(paragraph, text):
    """在段落中插入内容，并识别 **加粗** 区段。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for seg in parts:
        if seg.startswith("**") and seg.endswith("**"):
            inner = seg[2:-2]
            try:
                paragraph.add_run(inner, style="强调")
            except KeyError:
                run = paragraph.add_run(inner)
                run.bold = True
        else:
            paragraph.add_run(seg)


def add_codeblock(doc, lang, content):
    """添加代码块"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "TableGrid"
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]

    try:
        p.style = "CodeBlock"
    except KeyError:
        pass

    run = p.add_run(content)

    if hasattr(run, "_element"):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "Consolas")

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_heading(doc, level, text):
    """添加标题"""
    style_map = {
        1: "标题 1",
        2: "标题 2",
        3: "标题 3",
        4: "标题 4",
        5: "标题 5",
    }
    style_name = style_map.get(level, "正文")

    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        print(f"警告：样式 '{style_name}' 不存在，使用默认样式。", file=sys.stderr)
        p = doc.add_paragraph()

    p.add_run(text)


def add_list_item(doc, level, text, num_id):
    """添加列表项"""
    try:
        p = doc.add_paragraph(style="正文")
    except KeyError:
        p = doc.add_paragraph()

    apply_list_numbering(p, level, num_id)
    apply_bold_runs(p, text)


def add_paragraph(doc, text):
    """添加普通段落"""
    try:
        p = doc.add_paragraph(style="正文")
    except KeyError:
        p = doc.add_paragraph()
    apply_bold_runs(p, text)


def add_table(doc, rows, cfg):
    """添加表格"""
    custom_cfg = next(
        (ts for ts in cfg.get("table_styles", []) if ts["name"] == "自定义表格"),
        None,
    )

    table = doc.add_table(rows=len(rows), cols=len(rows[0]), style="TableGrid")

    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = value
            p = cell.paragraphs[0]
            run = p.runs[0]

            if custom_cfg:
                font_cfg = custom_cfg["font"]
                run.font.name = font_cfg["name"]
                run.font.size = Pt(font_cfg["size"])
                run.font.color.rgb = rgb(font_cfg["color"])

                if hasattr(run, "_element"):
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn("w:eastAsia"), font_cfg["name"])

                pf_cfg = custom_cfg.get("paragraph_format", {})
                p.paragraph_format.space_after = Pt(pf_cfg.get("space_after", 0))
                p.paragraph_format.line_spacing = pf_cfg.get("line_spacing", 1.0)

                if r == 0 and "header_row" in custom_cfg:
                    hdr = custom_cfg["header_row"]
                    run.font.bold = hdr["font"].get("bold", True)

                    if "fill_color" in hdr:
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:fill"), hdr["fill_color"])
                        shd.set(qn("w:val"), "clear")
                        tcPr.append(shd)

                    if "paragraph_format" in hdr:
                        align_map = {
                            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                        }
                        p.paragraph_format.alignment = align_map.get(
                            hdr["paragraph_format"].get("alignment"),
                            WD_ALIGN_PARAGRAPH.LEFT,
                        )
